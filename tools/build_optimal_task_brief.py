from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "artifacts" / "optimal-task-brief-reference.docx"

NAVY = "17324D"
BLUE = "2E74B5"
MID_BLUE = "4F81BD"
PALE_BLUE = "E8F0F7"
PALE_GRAY = "F3F5F7"
MID_GRAY = "66717D"
INK = "202A33"
WHITE = "FFFFFF"
GOLD = "A66A00"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DEE5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Aptos", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def keep_with_next(paragraph, enabled=True):
    paragraph.paragraph_format.keep_with_next = enabled


def no_split(paragraph):
    paragraph.paragraph_format.widow_control = True


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run, size=9, color=MID_GRAY)


def add_custom_numbering(doc, num_id, abstract_id, fmt, text, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def add_list_item(doc, text, num_id=900, bold_lead=None, size=10.5):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.20
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size)
    return p


def add_label_line(doc, label, prompt, guidance=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.18
    keep_with_next(p)
    r = p.add_run(label)
    set_run_font(r, color=NAVY, bold=True)
    r = p.add_run(f"  {prompt}")
    set_run_font(r, color=INK)
    if guidance:
        r = p.add_run(f"  {guidance}")
        set_run_font(r, size=9, color=MID_GRAY, italic=True)
    return p


def add_response_box(doc, height_lines=2, placeholder="[Type here]"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="CAD4DE", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(placeholder)
    set_run_font(run, size=9.5, color="77838F", italic=True)
    for _ in range(height_lines - 1):
        p.add_run("\n")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}  ")
    set_run_font(r, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_instruction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=MID_GRAY, italic=True)
    return p


def add_module_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2700, 6660], 120)
    set_table_borders(table)
    headers = ("Use when", "Add these details")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for index, (use_when, details) in enumerate(rows):
        cells = table.add_row().cells
        if index % 2:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for cell, text in zip(cells, (use_when, details)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(text)
            set_run_font(r, size=8.9, color=INK, bold=(cell is cells[0]))
    set_table_geometry(table, [2700, 6660], 120)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 13.5, MID_GRAY, 0, 18),
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 14, 6),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("OPTIMAL TASK BRIEF  /  REUSABLE DIRECTIVE")
    set_run_font(r, size=8.5, color=MID_GRAY, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PAGE ")
    set_run_font(r, size=8.5, color=MID_GRAY)
    add_field(p, "PAGE")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    add_custom_numbering(doc, 900, 900, "bullet", "•")
    add_custom_numbering(doc, 901, 901, "decimal", "%1.")
    add_custom_numbering(doc, 902, 902, "bullet", "•")

    # Page 1: editorial opening
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REUSABLE DIRECTIVE")
    set_run_font(r, size=9.5, color=GOLD, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Optimal Task Brief")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("A practical template for giving Codex the context, authority, constraints, and finish line needed to deliver excellent work.")

    add_callout(
        doc,
        "THE RULE",
        "Fill the Essentials every time. Add only the optional modules that materially change the work. Delete guidance text before sending if you want a cleaner brief.",
        fill="FFF7E8",
        accent=GOLD,
    )
    add_heading(doc, "How to use this template", 1)
    add_list_item(doc, "Start with the outcome: describe the finished result, not just the activity.", 901)
    add_list_item(doc, "Supply the context and source material that would change a good decision.", 901)
    add_list_item(doc, "State boundaries, permissions, and irreversible actions explicitly.", 901)
    add_list_item(doc, "Define how the result will be checked and what “done” means.", 901)
    add_list_item(doc, "For quick tasks, send only the one-minute brief below; omit the rest.", 901)

    add_heading(doc, "One-minute brief", 1)
    quick = doc.add_table(rows=5, cols=2)
    quick.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(quick, [1900, 7460], 120)
    set_table_borders(quick)
    quick_rows = [
        ("Outcome", "[What should exist or be true when finished?]"),
        ("Context", "[Why this matters; audience; current state.]"),
        ("Inputs", "[Files, links, examples, data, or prior work to use.]"),
        ("Constraints", "[Must / must not; deadline; format; scope; risk.]"),
        ("Done when", "[Tests, review criteria, and expected handoff.]"),
    ]
    for idx, (label, prompt) in enumerate(quick_rows):
        left, right = quick.rows[idx].cells
        set_cell_shading(left, NAVY)
        if idx % 2:
            set_cell_shading(right, "F8FAFC")
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(prompt)
        set_run_font(r, size=9.5, color=MID_GRAY, italic=True)
    set_table_geometry(quick, [1900, 7460], 120)

    # Page 2: essentials
    doc.add_page_break()
    add_heading(doc, "1. Essentials", 1)
    add_instruction(doc, "Complete this section for every task. Be concrete; short answers are enough.")
    add_label_line(doc, "Task name", "[Short, recognizable label]")
    add_response_box(doc, 1)
    add_label_line(doc, "Desired outcome", "[Describe the final state or artifact.]",
                   "Example: “A tested Python CLI that converts Suite2p outputs into tidy CSV files.”")
    add_response_box(doc, 2)
    add_label_line(doc, "Why this matters", "[Decision, workflow, user need, or problem this supports.]")
    add_response_box(doc, 2)
    add_label_line(doc, "Audience / user", "[Who will read, run, approve, or maintain the result?]")
    add_response_box(doc, 1)
    add_label_line(doc, "Current state", "[What exists now? What has already been tried?]")
    add_response_box(doc, 2)
    add_label_line(doc, "Inputs and sources", "[List or attach files, folders, links, examples, datasets, and source-of-truth material.]")
    add_response_box(doc, 2)
    add_callout(
        doc,
        "SOURCE PRIORITY",
        "If sources conflict, identify the authority order. Example: “The attached protocol overrides README text; existing tests override comments.”",
    )
    add_label_line(doc, "Deliverable(s)", "[Exact artifact, answer, change, or action expected; include format and location.]")
    add_response_box(doc, 2)
    add_label_line(doc, "Out of scope", "[What should remain untouched or explicitly not be solved?]")
    add_response_box(doc, 2)

    # Page 3: execution contract
    execution_heading = add_heading(doc, "2. Execution contract", 1)
    execution_heading.paragraph_format.page_break_before = False
    add_instruction(doc, "Use this section to set autonomy, risk, quality, and communication expectations.")
    add_label_line(doc, "Success criteria", "[Observable checks that distinguish complete from merely attempted.]")
    add_response_box(doc, 2)
    add_label_line(doc, "Constraints and invariants", "[Must keep; must avoid; compatibility; style; performance; compliance; budget.]")
    add_response_box(doc, 2)
    add_label_line(doc, "Authority and permissions", "[Actions allowed without asking; actions requiring approval.]")
    add_response_box(doc, 1)
    add_callout(
        doc,
        "GOOD DEFAULT",
        "Allow read-only inspection and reversible in-scope edits. Require confirmation for destructive actions, external messages, purchases, publication, credential use, or scope expansion.",
        fill=PALE_GRAY,
        accent=MID_BLUE,
    )
    add_label_line(doc, "Decision policy", "[When details are missing, should Codex infer, present options, or pause?]")
    add_response_box(doc, 1)
    add_label_line(doc, "Verification", "[Tests, renders, calculations, cross-checks, reviewers, or acceptance procedure.]")
    add_response_box(doc, 1)
    add_label_line(doc, "Communication", "[Update frequency, desired explanation depth, and when to escalate blockers.]")
    add_response_box(doc, 1)
    add_label_line(doc, "Final handoff", "[What the closing response must contain: links, summary, risks, commands, next steps.]")
    add_response_box(doc, 1)

    # Page 4: optional modules
    modules_heading = add_heading(doc, "3. Optional task modules", 1)
    modules_heading.paragraph_format.page_break_before = True
    add_instruction(doc, "Copy the relevant prompts into your brief. Skip modules that do not affect the result.")
    module_rows = [
        ("Code / debugging", "Repository and branch; runtime; target behavior; reproduction steps; logs; compatibility; files not to touch; tests and performance limits."),
        ("Research / analysis", "Decision to support; date range; geography; acceptable sources; freshness; evidence standard; citation format; assumptions; comparison criteria."),
        ("Documents / slides", "Audience; purpose; source material; length; brand or visual reference; structure; tone; editable output format; render and layout requirements."),
        ("Data / spreadsheets", "Input schema; units; missing-value rules; transformations; formulas versus values; validation totals; charts; output schema; privacy constraints."),
        ("Images / design", "Use context; dimensions; style references; composition; required and forbidden elements; text accuracy; brand colors; file format; variants."),
        ("Browser / UI work", "Target URL or local app; account/session assumptions; exact user flow; browsers/viewports; permitted submissions; screenshot and evidence needs."),
        ("Automation / monitoring", "Trigger; cadence; timezone; inputs; action; notification target; deduplication; failure handling; stop condition; archive policy."),
        ("External systems", "Named service; account/workspace; records in scope; read versus write authority; recipients; draft versus send; audit trail; rollback plan."),
        ("High-risk domain", "Relevant jurisdiction or policy; professional-review requirement; prohibited actions; uncertainty threshold; authoritative sources; explicit approval gates."),
    ]
    add_module_table(doc, module_rows)
    add_heading(doc, "Useful preference switches", 2)
    switches = [
        "Optimize for: [speed / correctness / maintainability / visual polish / cost / minimal change].",
        "Response style: [concise / explanatory / expert-level / tutorial].",
        "Tradeoffs: [choose the best default / show 2–3 options / preserve current approach].",
        "Progress: [work silently / send milestone updates / report each meaningful decision].",
        "Uncertainty: [state assumptions and continue / verify first / stop when confidence is below ___].",
    ]
    for item in switches:
        add_list_item(doc, item, 900, size=9.5)

    # Page 5: paste-ready master prompt
    master_heading = add_heading(doc, "4. Paste-ready master prompt", 1)
    master_heading.paragraph_format.page_break_before = True
    add_instruction(doc, "Replace bracketed text and remove unused lines. This is the shortest complete version of the template.")
    master_text = (
        "TASK\n"
        "[Name the task and describe the finished outcome.]\n\n"
        "CONTEXT\n"
        "[Explain why it matters, who it serves, the current state, and relevant history.]\n\n"
        "INPUTS\n"
        "[Attach or list files, links, examples, data, and sources. State source priority if they conflict.]\n\n"
        "DELIVERABLES\n"
        "[Specify artifacts or actions, formats, locations, and expected final handoff.]\n\n"
        "SCOPE AND CONSTRAINTS\n"
        "[State what is in scope, out of scope, must remain unchanged, and any deadline, compatibility, style, cost, privacy, or policy limits.]\n\n"
        "AUTHORITY\n"
        "[State what you may do without asking and what requires approval.]\n\n"
        "WORKING POLICY\n"
        "[State whether to infer reasonable details, offer options, or pause. Tell me how often to update you.]\n\n"
        "QUALITY BAR AND VERIFICATION\n"
        "[Define success criteria, required tests or checks, and what “done” means.]\n\n"
        "OPTIONAL TASK-SPECIFIC DETAILS\n"
        "[Add only the relevant module details from page 4.]\n\n"
        "Please inspect the available context first, make reasonable in-scope assumptions explicit, complete and verify the work, and finish with a concise handoff covering the outcome, changed artifacts, verification performed, remaining risks, and the most useful next step."
    )
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color=NAVY, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    for i, line in enumerate(master_text.split("\n")):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        is_label = line in {
            "TASK", "CONTEXT", "INPUTS", "DELIVERABLES", "SCOPE AND CONSTRAINTS",
            "AUTHORITY", "WORKING POLICY", "QUALITY BAR AND VERIFICATION",
            "OPTIONAL TASK-SPECIFIC DETAILS",
        }
        set_run_font(r, size=9.2, color=NAVY if is_label else INK, bold=is_label, italic=line.startswith("["))

    add_heading(doc, "Pre-send check", 1)
    checks = [
        "The finished outcome is concrete and observable.",
        "All necessary inputs are attached or locatable.",
        "Scope, protected areas, and approval gates are explicit.",
        "The quality bar and verification method are stated.",
        "Optional detail is included only where it changes execution.",
    ]
    for item in checks:
        add_list_item(doc, item, 902)

    # Core properties and save
    doc.core_properties.title = "Optimal Task Brief"
    doc.core_properties.subject = "Reusable directive template for briefing Codex effectively"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "task brief, prompt template, directive, Codex"
    doc.settings.element.set(qn("w:updateFields"), "true")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
